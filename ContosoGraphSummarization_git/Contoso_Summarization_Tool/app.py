import os  
import requests  
import asyncio  
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify  
from flask import Flask, request, jsonify, render_template_string, send_file
from markupsafe import Markup  # Updated import  
import markdown2 as markdown  # Updated import  
from graphrag_chat import graphchat  
from generate_graphrag_index import create_chat_index
from generate_summary import langraph
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import io
from docx import Document
import pandas as pd
  
app = Flask(__name__)  
app.secret_key = 'your_secret_key'  # Required for flashing messages  

# Configure the database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///books.db'  
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  
db = SQLAlchemy(app) 

class Book(db.Model):  
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)  
    title = db.Column(db.String(255), nullable=False)  
    author = db.Column(db.String(255), nullable=False)  
    synopsis = db.Column(db.Text, nullable=False)  
    characters = db.Column(db.Text, nullable=False)
    locations = db.Column(db.Text, nullable=False)
    readingage = db.Column(db.Text, nullable=False)
    audience = db.Column(db.Text, nullable=False)
    tone = db.Column(db.Text, nullable=False)
    themes = db.Column(db.Text, nullable=False)
    genres = db.Column(db.Text, nullable=False)
    readingguides = db.Column(db.Text, nullable=False)
    teachingguides = db.Column(db.Text, nullable=False)
    bookindex = db.Column(db.Text, nullable=False)
    single_tagline = db.Column(db.Text, nullable=False)
    full_plot_summary = db.Column(db.Text, nullable=False)
    update_date = db.Column(db.DateTime, nullable=False)
    
  
# Create the database and tables  
with app.app_context():  
    db.create_all()  

  
uploads_dir = os.path.join(app.root_path, 'uploads')  
  
# Ensure uploads directory exists  
os.makedirs(uploads_dir, exist_ok=True)  
  


@app.route('/', methods=['GET', 'POST'])  
def index():  
    selected_book = None  

    uploaded_books = os.listdir(uploads_dir)  # List of uploaded books  
  
    if request.method == 'POST':  
        if 'book' in request.files:  
            uploaded_file = request.files['book']  
            if uploaded_file and uploaded_file.filename != '':  
                file_path = os.path.join(uploads_dir, uploaded_file.filename)  
  
                # Save the uploaded file  
                uploaded_file.save(file_path)  
                flash('Book uploaded successfully!')  
  
                # Redirect to the index route to show the uploaded books  
                return redirect(url_for('index'))  
  
    return render_template('index.html', uploaded_books=uploaded_books, selected_book=selected_book)#, summary=summary)  



@app.route('/upload', methods=['POST'])  
def upload():  
    if 'book' in request.files:  
        uploaded_file = request.files['book']  
        if uploaded_file and uploaded_file.filename != '':  
            file_path = os.path.join(uploads_dir, uploaded_file.filename)  
  
            # Save the uploaded file  
            uploaded_file.save(file_path)  
            flash('Book uploaded successfully!')  
  
    # After uploading, get the list of uploaded books and return it as JSON  
    uploaded_books = os.listdir(uploads_dir)  
    return jsonify({'uploaded_books': uploaded_books})  


  
@app.route('/overwrite/<filename>', methods=['POST'])  
def overwrite(filename):  
    uploaded_file = request.files['book']  
    if uploaded_file:  
        file_path = os.path.join(uploads_dir, filename)  
        uploaded_file.save(file_path)  
        flash('Book overwritten successfully!')  
    return redirect(url_for('index'))  
 

# Function to add the book to the database  
def add_book_to_db(book_name, data):  
    new_book = Book(  
        filename=book_name,  # Use book_name from the request as bookname  
        title=data['title'],  
        author=data['author'],  
        synopsis=data['synopsis'],  
        characters=data['characters'],  
        locations=data['locations'],  
        readingage=data['readingage'],  
        audience=data['audience'],  
        tone=data['tone'],  
        themes=data['themes'],  
        genres=data['genres'],  
        readingguides=data['readingguides'],  
        teachingguides=data['teachingguides'],  
        bookindex=data['bookindex'],  
        single_tagline=data['single_tagline'],  
        full_plot_summary=data['full_plot_summary'],  
        update_date=datetime.now()  # Set current time as update date  
    )  
    db.session.add(new_book)  
    db.session.commit()
    print(f"Book added to database: {book_name}")
  


@app.route('/fetch_summary', methods=['GET'])  
async def fetch_summary():  
    book_name = request.args.get('book')  # Get book name from query parameters  
    print(f"Book name in Fetch Summary: {book_name}")  
    
    if not book_name:
        flash('No book specified.')  
        return jsonify({'error': 'No book specified.'}), 400  
  
    # Check if the book exists in the database  
    book = Book.query.filter_by(filename=book_name).first()  

         
    if book:  

        summary = {  
            'title': book.title,  
            'author': book.author,  
            'synopsis': book.synopsis,  
            'characters': book.characters,  
            'locations': book.locations,  
            'readingage': book.readingage,  
            'audience': book.audience,  
            'tone': book.tone,  
            'themes': book.themes,  
            'genres': book.genres,  
            'readingguides': markdown.markdown(book.readingguides),  
            'teachingguides': book.teachingguides,  
            'bookindex': markdown.markdown(book.bookindex),  
            'single_tagline': book.single_tagline,  
            'full_plot_summary': book.full_plot_summary,  
            'update_date': book.update_date.isoformat()  # Convert to ISO format for JSON  
        }  
        return jsonify(summary)  
  
    else:  
        data = await langraph(book_name)

        print(f"Data from fetch_summary: {data}")

        add_book_to_db(book_name, data)  # Pass book_name to add the book to the database  


        # Check if the book exists in the database  
        book = Book.query.filter_by(filename=book_name).first()  
        
        if book:  
            # If the book exists, return its summary  
            summary = {  
                'title': book.title,  
                'author': book.author,  
                'synopsis': book.synopsis,  
                'characters': book.characters,  
                'locations': book.locations,  
                'readingage': book.readingage,
                'audience': book.audience,  
                'tone': book.tone,  
                'themes': book.themes,  
                'genres': book.genres,  
                'readingguides': book.readingguides,  
                'teachingguides': book.teachingguides,  
                'bookindex': book.bookindex,  
                'single_tagline': book.single_tagline,  
                'full_plot_summary': book.full_plot_summary,  
                'update_date': book.update_date.isoformat()  # Convert to ISO format for JSON  
            }  

            return jsonify(summary)  
                 



import pandas as pd  
  
@app.route('/chat', methods=['POST'])  
def chat():  
    print("Inside chat")  
    data = request.json  
    question = data.get('question')  
    book_name = data.get('bookName')  
    searchType = data.get('searchType')
    print(f"Book name in chat function: {book_name}")  
    print(f"Search Type in chat function: {searchType}")
  
    # Implement your chat logic here  
    result = asyncio.run(graphchat(question, book_name, searchType))  # Assume graphchat returns a result object  
  
    # Render the answer as markdown  
    rendered_answer = markdown.markdown(result.response)  # Assuming result.response exists  
  
    # Initialize reports_json  
    reports_json = []  # Default to an empty list if no reports are available  
  
    # Convert context_data.reports DataFrame to a JSON-serializable format  
    #if hasattr(result.context_data, 'reports') and isinstance(result.context_data['reports'], pd.DataFrame):  
    reports_json = result.context_data['reports'].to_dict(orient='records')  # Convert to list of dictionaries  
    #print(f"Reports JSON: {reports_json}")
  
    response_details = {  
    'completion_time': result.completion_time,  # Assuming these attributes exist  
    'llm_calls': result.llm_calls,  
    'prompt_tokens': result.prompt_tokens,  
    } 

    #print(f"Response Details in app.py: {response_details}")
    
    # Return both the rendered answer and any additional context data  
    return jsonify({  
        'answer': rendered_answer,  
        'context_data': {  
            'reports': reports_json  # Use the converted reports  
        },
        'response_details': response_details
    })




@app.route('/create_index', methods=['POST'])  
def create_index():  
    book_name = request.form['book_name']  
    print(f"Book name in chat create_index: {book_name}")

    create_chat_index(book_name)
    return jsonify({"status": "success"})  

# export book

@app.route('/export_summary/<book_name>', methods=['GET'])  
def export_summary(book_name):  
    book = Book.query.filter_by(filename=book_name).first()    
    if book:  
        # Create a new Document  
        doc = Document()  
  
        # Add content to the document  
        doc.add_heading('Book Summary', level=1)  
        doc.add_paragraph(f'Title: {book.title}')  
        doc.add_paragraph(f'Author: {book.author}')  
        doc.add_paragraph(f'Synopsis: {book.synopsis}')  
        doc.add_paragraph(f'Characters: {book.characters}')  
        doc.add_paragraph(f'Locations: {book.locations}')  
        doc.add_paragraph(f'Reading Age: {book.readingage}')  
        doc.add_paragraph(f'Audience: {book.audience}')  
        doc.add_paragraph(f'Tone: {book.tone}')  
        doc.add_paragraph(f'Themes: {book.themes}')  
        doc.add_paragraph(f'Genres: {book.genres}')  
        doc.add_paragraph(f'Reading Guides: {book.readingguides}')  
        doc.add_paragraph(f'Teaching Guides: {book.teachingguides}')  
        doc.add_paragraph(f'Book Index: {book.bookindex}')  
        doc.add_paragraph(f'Single Tagline: {book.single_tagline}')  
        doc.add_paragraph(f'Full Plot Summary: {book.full_plot_summary}')  
        doc.add_paragraph(f'Update Date: {book.update_date.isoformat()}')  
  
        # Save the document to a BytesIO object  
        doc_io = io.BytesIO()  
        doc.save(doc_io)  
        doc_io.seek(0)  
  
        # Send the file as a response  
        return send_file(doc_io, as_attachment=True, download_name=f"{book.title}_summary.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')  
    else:  
        return jsonify({'error': 'Book not found'}), 404 

if __name__ == '__main__':  
    with app.app_context():  
        db.create_all()  # Create tables if they don't exist  
    app.run(debug=True) 
