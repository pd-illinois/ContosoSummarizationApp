from langchain_openai import AzureChatOpenAI
from dotenv import load_dotenv
from langchain.text_splitter import CharacterTextSplitter
import tiktoken
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain import hub
import operator
from typing import Annotated, List, Literal, TypedDict
from langchain.chains.combine_documents.reduce import (
    acollapse_docs,
    split_list_of_docs,
)
from langchain_core.documents import Document
from langgraph.constants import Send
from langgraph.graph import END, START, StateGraph
import asyncio
import json
import os  
from docx import Document  
from dotenv import load_dotenv


import os

async def langraph (book_name):
    # Import Azure OpenAI
    from langchain_openai import AzureChatOpenAI

    load_dotenv("./config/.env")
    llm_model = os.getenv("llm_model")
    llm_heavy_model = os.getenv("llm_heavy_model")  
    api_version = os.getenv("api_version")
    api_base = os.getenv("api_base")
    api_key = os.getenv("api_key")

    llm = AzureChatOpenAI(
        deployment_name=llm_model,
        api_version=api_version,
        azure_endpoint=api_base,
        api_key=api_key,
        streaming=False
    )

    llm_heavy = AzureChatOpenAI(
        deployment_name=llm_heavy_model,
        api_version=api_version,
        azure_endpoint=api_base,
        api_key=api_key,
        streaming=False
    )

    from docx import Document
    from langchain.text_splitter import CharacterTextSplitter

    def read_docx(file_path):
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)


    def split_into_chunks_with_langchain(bookcontents, number_of_chunks, overlap=0):
        # Calculate the chunk size
        chunk_size = len(bookcontents) // number_of_chunks
        
        # Initialize the text splitter
        text_splitter = CharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap
        )
        
        # Split the text into chunks
        chunks = text_splitter.split_text(bookcontents)
        
        return chunks

    #Import encoding for gpt4-o and get chunks
    import tiktoken
    uploads_dir = os.getenv("uploads_dir")
    book_path = os.path.join(uploads_dir, book_name)
    print(f'book path is {book_path}')

    text = read_docx(book_path)

    chunksize = 5000
    overlap = 1000

    encoding = tiktoken.encoding_for_model("gpt-4o")
    gpt4encoded = encoding.encode(text)

    number_of_chunks = len(gpt4encoded) // chunksize + 1
    chunks = split_into_chunks_with_langchain(text, number_of_chunks, overlap)

    len(chunks),number_of_chunks

    from langchain_core.output_parsers import StrOutputParser
    from langchain_core.prompts import ChatPromptTemplate
    from langchain import hub

    index_prompt = """

    Index
    An alphabetical list of important topics, names, places, and key terms mentioned within the book, count the number of times they show. 
    Categorize the index entries as follows:
    - Characters
    - Locations
    - Events
    - Objects

    Example 
    Alice : Number of Appearences - 22

    """

    booksynopsis_prompt = """

    Create a concise and engaging synopsis of a novel. Follow these guidelines:

    Keep it brief: Condense the manuscript into a couple of pages by expanding the book blurb into a full synopsis. Aim for succinct writing.
    Mention only the most important details: Focus on key characters and essential plot points. Include only characters crucial to the storyline.
    Keep the writing style simple: Use clear and straightforward language. Avoid elaborate prose to maximize the limited space.
    Add emotion: Convey the story's emotions vividly using strong adjectives, adverbs, and verbs. Ensure the synopsis tells a condensed yet engaging story.
    Reveal the ending: Unlike a book blurb, the synopsis must reveal the ending to show the narrative arc and marketability.
    Use these guidelines to produce a high-quality synopsis that effectively represents the novel and entices editors.
    However the output should only be a synopsis, no separate headers for characters, locations, events or objects.
    If you do not know who the writer is, do not make any assumptions and do not credit it to unknown writer. 

    Here is an example:
    Pride and Prejudice follows the life of Elizabeth Bennet, one of five sisters in a respectable but not wealthy family in 19th-century England. The novel explores themes of love, social standing, and personal growth. Elizabeth's 
    initial prejudice against the wealthy and aloof Mr. Darcy evolves as she learns more about his true character. Through a series of misunderstandings, social events, and personal revelations, Elizabeth and Darcy overcome their pride and prejudices, ultimately finding love and mutual respect. 
    The novel concludes with their marriage, symbolizing the triumph of true affection and understanding over societal expectations and superficial judgments.

    """

    bookdetails_prompt = """

    Title of the book: Get the name of the book.
    Author of the book: Get the name of the author, it will not say Author and then the name. It should follow directly after the first mention of the titla. 
    Names of Characters and Character Descriptions: Detailed information about the characters. Mention relationships to other characters. Include all physical descriptions.
    Locations/Setting: Description of the book's setting.
    Reading Group Age: Target age group for the book + a reasoning why it is suitable for that age group. Consider the complexity of the language, themes, and content.
    Potential Audience/Demographics: Target audience and demographic information.
    Tone: The attitude or feeling conveyed by the author through word choice and language. Examples include formal, objective, whimsical, persuasive, and conversational.
    Themes: The central idea, message, or lesson of the book. Examples include unrequited love, friendship, true heroism, and fear.
    Genres/Categories: Examples include science fiction, crime, history, memoir, romance, action adventure/thriller, mystery, horror/dystopian, mythology, philosophy, alternative belief, educational texts, and reference. 

    Format your output exactly as follows, as they will be part of a larger JSON object.
    You MUST NOT add curly braces, and you MUST end each key/value with a comma.
    All linefeeds within the answers should be either escaped or in markdown format. The answers themselves between the quotes should be a single line.

    "title": "<fill in your answer, you must use markdown for readability>",
    "author": "<fill in your answer here, you must use markdown for readability>",
    "characters": "<fill in your answer, you must use markdown for readability>",
    "locations": "<fill in your answer here, you must use markdown for readability>",
    "readingage": "<fill in your answer here, you must use markdown for readability>",
    "audience": "<fill in your answer here, you must use markdown for readability>",
    "tone": "<fill in your answer here, you must use markdown for readability>",
    "themes": "<fill in your answer here, you must use markdown for readability>",
    "genres": "<fill in your answer here, you must use markdown for readability>",

    """

    readersguide_prompt = """

    Readers and Teachers Guides
    Reading Guides: Questions and prompts created by teachers to help students comprehend the main points of the reading and understand the structure of the text. The example questions must be specific to the characters, motivations and setting of the book.
    Teacher Created Reading Guides: Guides to help students navigate challenging reading material. List important themes, symbols, and motifs in the book.

    Format your output exactly as follows, as they will be part of a larger JSON object.
    You MUST NOT add curly braces, and you MUST end each key/value with a comma.
    All linefeeds within the answers should be either escaped or in markdown format. The answers themselves between the quotes should be a single line.

    "readingguides": "<fill in your answer here, you must use markdown for readability>",
    "teachingguides": "<fill in your answer here, you must use markdown for readability>",

    A real example of a reading guide could be:

    "readingguides": "<br> 1. How does Alice's perception of her relationship with Nolan change throughout the story? <br>- 2. Discuss the role of family dynamics in the Altman family and how they affect Nolan's character."

    """

    information_prompt = """Index: Create an alphabetical list of important topics, names, places, and key terms mentioned within the book, along with corresponding page numbers. Categorize entries as Characters, Locations, Events, and Objects. Book Synopsis: Create a concise and engaging synopsis of the novel. Follow these guidelines:
    Keep it brief: Condense the manuscript into a couple of pages.
    Mention only the most important details: Focus on key characters and essential plot points.
    Keep the writing style simple: Use clear and straightforward language.
    Add emotion: Convey the story's emotions vividly.
    Reveal the ending: Show the narrative arc and marketability.
    Do not use separate headers for characters, locations, events, or objects.
    Do not make assumptions about the writer if unknown.
    Book Details: Provide detailed information about the book:
    Names of Characters and Character Descriptions
    Locations/Setting
    Reading Group Age
    Potential Audience/Demographics
    Tone
    Themes
    Genres/Categories
    Format your output exactly as follows, as they will be part of a larger JSON object. Do not add curly braces, and end each key/value with a comma. Use markdown for readability.
    Readers and Teachers Guides: Create reading and teaching guides:
    Reading Guides: Questions and prompts to help students comprehend the main points and understand the structure of the text.
    Teacher Created Reading Guides: Guides to help students navigate challenging reading material, listing important themes, symbols, and motifs.
    Format your output exactly as follows, as they will be part of a larger JSON object. Do not add curly braces, and end each key/value with a comma. Use markdown for readability."

    """

    # map_prompt = hub.pull("rlm/map-prompt")
    map_prompt = ChatPromptTemplate.from_template(F"""
            You will be provided with a section of a larger book. Your task is to write a plot summary of the text. You cannot use any information from outside the text.
            The text should reflect the tone and style of the original book, and should contain as much detailed information about the plot, characters, locations and events as possible.
            
            You will only output the summary, but keep in mind that we ultimately want to use the summaries to extract the following information:
            
            {information_prompt}
            
            If you find out with relative certainty what the title of the book is, or who wrote it, prominently feature it in the summary
            Here is the section of the book:
            {{doc}}""")
    map_chain = map_prompt | llm | StrOutputParser()

    final_bookdetails_prompt = ChatPromptTemplate.from_template(f"""
            You will be provided with a set of partial summaries of a larger book. Based on these summaries, your task is to extract the following information:
            
            {bookdetails_prompt}
            
            Here are the summaries:
            {{doc_summaries}}""")
    final_bookdetails_chain = final_bookdetails_prompt | llm_heavy | StrOutputParser()

    final_index_scan_prompt = ChatPromptTemplate.from_template("""
            You will be provided with a set of partial summaries of a larger book. Based on these summaries, your task is to extract the following information:
            
            You will list out all of the important characters, locations, events and objects that are important the plot.
            
            Here are the summaries:
            {doc_summaries}""")
    final_index_scan_chain = final_index_scan_prompt | llm | StrOutputParser()

    final_index_prompt = ChatPromptTemplate.from_template("""
            You will be provided with many indexes from different parts of the book.
            Consolidate the indexes into a single alphabetical index, categorized as follows:
            
            - Characters
            - Locations
            - Events
            - Objects
            
            Here is an example of what the Character index could look like:
            
    ### Characters: Frequency\\n- Bar Comas: High\\n- Burroughs, Edgar Rice: Low\\n- Carter, Captain Jack: : Meduim

            The actual characters will depend on the indexes you are provided, do not use "Bar Comas", "Edgar Rice Burroughs" or "Captain Jack Carter" in the final index.
            Make sure to add the number of instances together by entity from all indexes. So if Alice Appears 4 times in one index and 10 times in another index, the result would be Alice : Number of instances 14. Based on the number give it a rating for how frequent it appears High, Meduim, Low. 
            Order entities by rating, High first, meduim next and low last.
            Do not add any header "Consolidated Index" or similar.
            Do not use the > or < symbols.
            Do not add any additional writing that explains the index.
            You must use markdown.
            
            {index_entries}
            """)
    final_index_chain = final_index_prompt | llm_heavy | StrOutputParser()

    map_index_prompt = ChatPromptTemplate.from_template(f"""
            You will be provided part of a book. You will look out for mentions of the following characters, locations, events and objects:
            
            {{index_entities}}
            
            You will create the following {index_prompt}
            Look for references to the chapters in each section given. 
            Here is the section of the book:
            {{doc}}""")
    map_index_chain = map_index_prompt | llm | StrOutputParser()

    final_readersguide_prompt = ChatPromptTemplate.from_template(f"""
            You will be provided with a set of partial summaries of a larger book. Based on these summaries, your task is to extract the following information:
            
            {readersguide_prompt}
            
            Here are the summaries:
            {{doc_summaries}}""")
    final_readersguide_chain = final_readersguide_prompt | llm | StrOutputParser()

    final_booksynopsis_prompt = ChatPromptTemplate.from_template(f"""
            You will be provided with a set of partial summaries of a larger book. 
            Based on these summaries, your task is write a full synopsis, following this guidance.
            Output in plain multiline text format.
            
            {booksynopsis_prompt}
            
            
            Here are the summaries:
            {{doc_summaries}}""")
    final_booksynopsis_chain = final_booksynopsis_prompt | llm_heavy | StrOutputParser()

    final_single_tagline_prompt = ChatPromptTemplate.from_template("""
            You will be provided with a set of partial summaries of a larger book. 
            Based on these summaries, your task is write a catchy, intriguing and interesting one-line description of the book.
            It is ok to make it a longer sentence to capture enough detail so that the line is specific to the characters, plot and setting of this book.
            The objective is to entice the audience to read the book.        
            
            Here are the summaries:
            {doc_summaries}""")
    final_single_tagline_chain = final_single_tagline_prompt | llm_heavy | StrOutputParser()

    final_full_plot_summary_prompt = ChatPromptTemplate.from_template("""
            You will be provided with a set of partial summaries of a larger book. 
            Write a neutral full plot summary that covers all of the important plot elements, characters, and events in the book in one cohesive narrative.
            The summary should be detailed and comprehensive, but not overly long, and should include the ending of the book.
            You must not use headers for characters, locations, events or objects, but rather create one condensed version of the plot.
            Ensure that the entire plot is covered in the summary, and the summary is written in a way that is clear, flowing and easy to read.
            Structure your summary in paragraphs to enhance readability.
            
            Here are the summaries:
            {doc_summaries}""")
    final_full_plot_summary_chain = final_full_plot_summary_prompt | llm_heavy | StrOutputParser()

    # reduce_prompt = hub.pull("rlm/reduce-prompt")
    reduce_prompt = ChatPromptTemplate.from_template(f"""
            You will be provided with a set of partial summaries of a larger book. Based on these summaries, your task is to extract the following information:
            
            {information_prompt}
            
            Here are the summaries:
            {{doc_summaries}}""")
    reduce_chain = reduce_prompt | llm | StrOutputParser()

    import operator
    from typing import Annotated, List, Literal, TypedDict

    from langchain.chains.combine_documents.reduce import (
        acollapse_docs,
        split_list_of_docs,
    )
    from langchain_core.documents import Document
    from langgraph.constants import Send
    from langgraph.graph import END, START, StateGraph

    token_max = chunksize + 2000

    def length_function(documents: List[Document]) -> int:
        """Get number of tokens for input contents."""
        return sum(llm.get_num_tokens(doc.page_content) for doc in documents)

    # This will be the overall state of the main graph.
    # It will contain the input document contents, corresponding
    # summaries, and a final summary.

    class OverallState(TypedDict):
        # Notice here we use the operator.add
        # This is because we want combine all the summaries we generate
        # from individual nodes back into one list - this is essentially
        # the "reduce" part
        contents: List[str]
        summaries: Annotated[list, operator.add]
        indexes: Annotated[list, operator.add]
        collapsed_summaries: List[Document]
        final_summary: str
        book_details: str
        index_entities: str
        index_final: str
        single_tagline: str
        full_plot_summary: str
        readers_guide: str
        synopsis: str
        


    # This will be the state of the node that we will "map" all
    # documents to in order to generate summaries
    class SummaryState(TypedDict):
        content: str


    # Here we generate a summary, given a document
    async def generate_summary(state: SummaryState):
        response = await map_chain.ainvoke(state["content"])
        return {"summaries": [response]}

    async def generate_index(state: OverallState):
        response = await map_index_chain.ainvoke({'doc' : state["content"], 'index_entities' : 'test'})
        return {"indexes": [response]}

    # Here we define the logic to map out over the documents
    # We will use this an edge in the graph
    def map_summaries(state: OverallState):
        # We will return a list of `Send` objects
        # Each `Send` object consists of the name of a node in the graph
        # as well as the state to send to that node
        return [
            Send("generate_summary", {"content": content}) for content in state["contents"]
        ]

    def map_indexes(state: OverallState):
        # We will return a list of `Send` objects
        # Each `Send` object consists of the name of a node in the graph
        # as well as the state to send to that node
        return [
            Send("generate_indexes",{"content": content}) for content in state["contents"]
        ]


    def collect_summaries(state: OverallState):
        # print(f"number of summaries : {len(state['summaries'])}")
        return {
            "collapsed_summaries": [Document(summary) for summary in state["summaries"]]
        }


    # Add node to collapse summaries
    async def collapse_summaries(state: OverallState):
        doc_lists = split_list_of_docs(
            state["collapsed_summaries"], length_function, token_max
        )
        results = []
        for doc_list in doc_lists:
            results.append(await acollapse_docs(doc_list, reduce_chain.ainvoke))

        return {"collapsed_summaries": results}

    # This represents a conditional edge in the graph that determines
    # if we should collapse the summaries or not
    def should_collapse(
        state: OverallState,
    ) -> Literal["collapse_summaries", "generate_final_summary"]:
        num_tokens = length_function(state["collapsed_summaries"])
        if num_tokens > token_max:
            return "collapse_summaries"
        else:
            return "generate_final_summary"


    # Here we will generate the final summary
    async def generate_book_details_summary(state: OverallState):
        response = await final_bookdetails_chain.ainvoke(state["collapsed_summaries"])
        return {"book_details": response}

    async def generate_index_scan(state: OverallState):
        response = await final_index_scan_chain.ainvoke(state["collapsed_summaries"])
        return {"index_entities": response}

    async def generate_synopsis(state: OverallState):
        response = await final_booksynopsis_chain.ainvoke(state["collapsed_summaries"])
        return {"synopsis": response}

    async def generate_readers_guide_summary(state: OverallState):
        response = await final_readersguide_chain.ainvoke(state["collapsed_summaries"])
        return {"readers_guide": response}

    async def generate_full_plot_summary(state: OverallState):
        response = await final_full_plot_summary_chain.ainvoke(state["collapsed_summaries"])
        return {"full_plot_summary": response}

    async def generate_single_tagline(state: OverallState):
        response = await final_single_tagline_chain.ainvoke(state["collapsed_summaries"])
        return {"single_tagline": response}

    async def generate_final_index(state: OverallState):
        response = await final_index_chain.ainvoke(state["indexes"])
        return {"index_final": response}

    def ensure_comma(text):
        """
        Ensures the last line of the given text ends with a comma, ignoring the last line if it's empty.
        
        Args:
            text (str): The input text.
            
        Returns:
            str: The modified text with the last line ending with a comma if it's not empty.
        """
        lines = text.split('\n')
        if lines and lines[-1].strip() == '':
            lines.pop()
        if lines and not lines[-1].strip().endswith(','):
            lines[-1] = lines[-1].strip() + ','
        return '\n'.join(lines)

    async def generate_final_summary(state: OverallState):
        full_plot_summary_escaped = state["full_plot_summary"].replace('"', '\\"').replace('\n', '\\n')
        single_tagline_escaped = state["single_tagline"].replace('"', '\\"').replace('\n', '\\n')
        synopsis_escaped = state["synopsis"].replace('"', '\\"').replace('\n', '\\n')
        index_escaped = state["index_final"].replace('"', '\\"').replace('\n', '\\n')
        book_details = ensure_comma(state["book_details"]).replace("\\'","'")
        readers_guide = ensure_comma(state["readers_guide"]).replace("\\'","'")
        response = f"""
            
        {{
            "synopsis": "{synopsis_escaped}", 
            {book_details} 
            {readers_guide}
            "bookindex": "{index_escaped}",
            "single_tagline": "{single_tagline_escaped}",
            "full_plot_summary": "{full_plot_summary_escaped}"
        }}
        
        """
        return {"final_summary": response}

    # Construct the graph
    # Nodes:
    graph = StateGraph(OverallState)
    graph.add_node("generate_summary", generate_summary)  # same as before
    graph.add_node("collect_summaries", collect_summaries)
    # graph.add_node("collapse_summaries", collapse_summaries)
    graph.add_node("generate_book_details_summary", generate_book_details_summary)
    graph.add_node("generate_index_scan", generate_index_scan)
    graph.add_node("generate_indexes", generate_index)
    graph.add_node("generate_index", generate_final_index)
    graph.add_node("generate_synopsis", generate_synopsis)

    graph.add_node("generate_readers_guide_summary", generate_readers_guide_summary)
    graph.add_node("generate_final_summary", generate_final_summary)

    graph.add_node("generate_single_tagline", generate_single_tagline)
    graph.add_node("generate_full_plot_summary", generate_full_plot_summary)


    # Edges:
    graph.add_conditional_edges(START, map_summaries, ["generate_summary"])
    graph.add_conditional_edges("generate_index_scan", map_indexes, ["generate_indexes"])

    graph.add_edge("generate_summary", "collect_summaries")

    # graph.add_conditional_edges("collect_summaries", should_collapse)
    # graph.add_conditional_edges("collapse_summaries", should_collapse)

    graph.add_edge("collect_summaries", "generate_synopsis")
    graph.add_edge("collect_summaries", "generate_book_details_summary")
    graph.add_edge("collect_summaries", "generate_index_scan")
    graph.add_edge("collect_summaries", "generate_single_tagline")
    graph.add_edge("collect_summaries", "generate_full_plot_summary")

    graph.add_edge("generate_indexes", "generate_index")
    graph.add_edge("generate_index", "generate_final_summary")
    graph.add_edge("collect_summaries", "generate_readers_guide_summary")
    graph.add_edge(["generate_single_tagline","generate_full_plot_summary","generate_synopsis","generate_book_details_summary","generate_index","generate_readers_guide_summary"], "generate_final_summary")
    # graph.add_edge("generate_index_scan", "generate_index")
    graph.add_edge("generate_final_summary", END)

    app = graph.compile()

    async for step in app.astream(
        {"contents": [doc for doc in chunks]},
    ):
        print(list(step.keys()))
        
    import json
    # Need to figure of use of \n and \\n
    book = json.loads(step['generate_final_summary']['final_summary'].replace("\\\\n","\\n"))
    #book = json.loads(step['generate_final_summary']['final_summary'].replace("\\n","\n"))
    print(book)
    return book